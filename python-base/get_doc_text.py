# _*_coding :utf-8 _*_
# @Time :2022/7/10 22:27
# @File : 测试
# @Project : python-base

import docx         # 下载python-docx模块

doc = docx.Document(r"C:\Users\yui\Documents\稻壳文档\学习四史.docx")
print(len(doc.paragraphs))              # 获取 总段落数


def gettext():                          # 读取文档的全部内容
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)

    return '\n'.join(fulltext)              # 用换行符进行拼接





